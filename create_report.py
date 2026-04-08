import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_agrisense_report(filename):
    doc = Document()
    
    # helper for styling
    def add_title(text, size=24):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        return p

    def add_center_bold(text, size=14):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        return p
        
    def add_heading(text, level=1):
        h = doc.add_heading(level=level)
        run = h.add_run(text)
        run.font.color.rgb = docx.shared.RGBColor(0,0,0)

    # Cover Page
    add_title("AGRISENSE: AN AI-POWERED AGRICULTURAL INTELLIGENCE PLATFORM")
    add_center_bold("A PROJECT REPORT\n", 16)
    
    add_center_bold("Submitted by", 14)
    names = "Suthendran K\nHemaswathi S\nVikas Kannan (9923008143)\nShanmugeshwar U (9923008134)\nGaushik S (9923008125)\nChriswin Raj (9923008112)"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(names).italic = True
    
    doc.add_paragraph()
    add_center_bold("in partial fulfillment for the award of the degree\nof\nBACHELOR OF TECHNOLOGY", 12)
    doc.add_paragraph()
    add_center_bold("DEPARTMENT OF INFORMATION TECHNOLOGY", 14)
    add_center_bold("KALASALINGAM ACADEMY OF RESEARCH AND EDUCATION\n(DEEMED TO BE UNIVERSITY)\nANAND NAGAR, KRISHNANKOIL – 626 126", 14)
    doc.add_page_break()

    # Declaration
    add_title("DECLARATION", 16)
    decl = ('I do hereby declare that this Project work entitled "AGRISENSE: AN AI-POWERED AGRICULTURAL INTELLIGENCE PLATFORM" '
            'submitted by us is the result of our original and independent research work carried out during for the project '
            'under the guidance of Kalasalingam University, Anand Nagar, Krishnankoil and it has not been submitted '
            'for the award of any Degree, Diploma, Associateship, Fellowship of any University or Institution.')
    doc.add_paragraph(decl)
    doc.add_paragraph("\nTeam Details:")
    doc.add_paragraph(names)
    doc.add_paragraph("\nDate: _______________")
    doc.add_page_break()

    # TABLE OF CONTENTS
    add_heading("TABLE OF CONTENTS")
    contents = [
        "1. Abstract",
        "2. Introduction (Problem Context and Need Identification)",
        "3. Literature Review and Gap Identification",
        "4. Problem Definition and Design Objectives",
        "5. Concept Generation and Selection",
        "6. System Design and Engineering Details",
        "7. Prototype Fabrication and Assembly (Software Implementation)",
        "8. Performance Testing and Evaluation",
        "9. Cost Analysis and Sustainability Assessment",
        "10. Discussion and Design Optimization",
        "11. Learning Outcomes and Team Reflection",
        "12. Conclusion and Future Scope",
        "13. References"
    ]
    for c in contents:
        doc.add_paragraph(c)
    doc.add_page_break()

    # 1. Abstract
    add_heading("1. Abstract")
    doc.add_paragraph("Farming has always been central to the Indian economy, but many farmers still make key decisions—like which crop to grow or how much yield to expect—based mostly on past experience rather than data. In this project, we present AgriSense, a web-based platform that brings together machine learning and generative AI to help farmers make better-informed choices. We use a Random Forest Classifier that considers soil and weather data (N, P, K, temperature, humidity, pH, and rainfall) to return the top three most suitable crops. For yield estimation, a Random Forest Regressor predicts the expected harvest based on area, rainfall, fertilizer, and pesticide usage. The platform also offers profit analysis and risk assessment modules. A standout feature is the integration of Google Gemini 2.5 Flash, which provides plain-language AI explanations for predictions. The complete system uses a FastAPI back-end and a React/Vite front-end, successfully achieving highly accurate predictive mapping for precision agriculture.")
    
    # 2. Introduction
    add_heading("2. Introduction (Problem Context and Need Identification)")
    doc.add_paragraph("Background: Agriculture is the lifeline of India, providing livelihoods to over 58% of the population. However, farmers continually face challenges with unpredictable weather, soil degradation, and market volatility.\n"
                      "Need Identification: Most farmers lack access to specialized, data-driven advice. When relying only on traditional knowledge, poor decisions regarding crop selection and input usage result in low yields and financial loss.\n"
                      "Objectives: To develop a unified web platform mapping soil/climate data to crop predictions, estimate yields, and translate numeric predictions into human-readable text via LLMs.")

    # 3. Literature Review
    add_heading("3. Literature Review and Gap Identification")
    doc.add_paragraph("Existing solutions for crop recommendation often use models like Naïve Bayes or SVM, but Random Forest consistently outperforms them due to its ensemble nature. However, most existing applications act as black-boxes—they tell the farmer what to plant without explaining why.\nGap: Current decision support systems lack natural language interpretability and do not combine yield, profit, and crop suitability into one unified tool accessible to non-technical users.")

    # 4. Problem Definition
    add_heading("4. Problem Definition and Design Objectives")
    doc.add_paragraph("Problem Statement: To design and develop an accurate, scalable, and interpretable agricultural intelligence software system that assists farmers in making optimal cultivation and financial decisions.\n"
                      "Design Criteria: \n"
                      "- Accuracy: Classification models must exceed 95% accuracy.\n"
                      "- Usability: The UI must be responsive and intuitive.\n"
                      "- Interpretability: Every model prediction must involve a contextual summary via Generative AI.")

    # 5. Concept Selection
    add_heading("5. Concept Generation and Selection")
    doc.add_paragraph("Multiple architectural concepts were evaluated:\n"
                      "1. Desktop App + Local Database (High performance, poor accessibility)\n"
                      "2. Monolithic Web App (Easy to deploy, harder to scale ML models)\n"
                      "3. Service-Oriented Web App (FastAPI + React) (Highly modular, scalable, excellent for ML integration)\n"
                      "Concept 3 was selected due to its separation of concerns and superior handling of REST API connections with external services like Google Gemini.")

    # 6. System Design
    add_heading("6. System Design and Engineering Details")
    doc.add_paragraph("6.1 Design Overview\nThe system uses a three-tier architecture: Data/Model Layer (scikit-learn, datasets, pickled models), Backend Layer (FastAPI, ML Inference, Gemini API connector), and Frontend Layer (React 18, Vite, Recharts).\n\n"
                      "6.2 Working Principle\nData inputs from the UI are passed as JSON to the API. The API processes categorical variables via a One-Hot-Encoder pipeline, runs Random Forest predictions, computes business logic (profits/risk functions), and constructs a dynamic text prompt sent to the Gemini API.\n\n"
                      "6.3 Technologies Used\nPython, FastAPI, Scikit-learn, React, Vite, Axios, Google Gemini SDK.")

    # 7. Prototype Fabrication
    add_heading("7. Software Implementation and Assembly")
    doc.add_paragraph("The development involved creating offline training scripts (train_classifier.py and train_regressor.py) that produced .pkl artifacts. The backend was organized using API routing files (/recommend, /yield, /compare) attached to Service classes. The frontend was developed using glassmorphism UI principles, managing global state for parameters and passing them to Recharts for immediate visual feedback.")

    # 8. Performance Testing
    add_heading("8. Performance Testing and Evaluation")
    doc.add_paragraph("We evaluated the system on publicly available datasets (2,200 samples for classification, 10,000+ for regression). \n"
                      "Results:\n"
                      "- Crop Recommendation: Random Forest achieved 99.32% accuracy.\n"
                      "- Yield Prediction: R-squared score of 0.9641.\n"
                      "Comparison: Random Forest outperformed Gradient Boosting (98.95%) and SVM (97.73%). The Gemini API returned contextually accurate and highly readable responses, rated qualitatively at 4.4/5.")

    # 9. Cost Analysis
    add_heading("9. Cost Analysis and Sustainability Assessment")
    doc.add_paragraph("Cost of Development: The platform relies entirely on open-source technologies (Python, React) and free-tier API services (Gemini), resulting in zero software licensing costs.\n"
                      "Sustainability: By predicting exact fertilizer and pesticide needs, AgriSense minimizes chemical runoff, directly promoting sustainable ecosystem preservation and soil health.")

    # 10. Discussion
    add_heading("10. Discussion and Design Optimization")
    doc.add_paragraph("The integration of Generative AI successfully bridged the gap between numeric data and actionable farming advice. However, the current risk heuristic is simple and relies only on rainfall/fertilizer anomalies. Furthermore, model predictions operate on static pickle files without live retraining capabilities.\n"
                      "Optimization: Implementing asynchronous API calls successfully reduced the UI hanging time when waiting for Gemini responses.")

    # 11. Learning Outcomes
    add_heading("11. Learning Outcomes and Team Reflection")
    doc.add_paragraph("Technical Skills: The team gained deep expertise in full-stack web development (FastAPI/React), building Machine Learning pipelines, and prompt engineering for LLMs.\n"
                      "Non-technical Skills: We learned agile project management, dividing tasks across the backend API, frontend UI, and model training efficiently.")

    # 12. Conclusion
    add_heading("12. Conclusion and Future Scope")
    doc.add_paragraph("AgriSense successfully demonstrates that classical machine learning can be vastly improved for end-users when paired with Generative AI. The dual-model pipeline is highly accurate for crop and yield prediction.\n"
                      "Future Scope: Extending the AI insight engine to support regional languages (Tamil, Hindi), integrating real-time IoT soil sensors for live data, and developing a native mobile application for low-bandwidth farming areas.")

    # 13. References
    add_heading("13. References")
    doc.add_paragraph("[1] K. G. Liakos et al., 'Machine learning in agriculture: A review,' Sensors, 2018.\n"
                      "[2] A. Kamilaris et al., 'Deep learning in agriculture: A survey,' Computers and Electronics in Agriculture, 2018.\n"
                      "[3] Google DeepMind, 'Gemini: A family of highly capable multimodal models,' arXiv preprint, 2024.")

    doc.save(filename)
    print(f"Report successfully saved to {filename}")

if __name__ == "__main__":
    create_agrisense_report(r"C:\Users\geshw\Agrisense\Agrisense_Final_Report.docx")
